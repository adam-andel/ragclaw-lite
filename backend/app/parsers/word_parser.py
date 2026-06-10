"""Word (.docx) document parser using python-docx."""

from pathlib import Path
from docx import Document as DocxDocument
from docx.oxml.ns import qn

from app.parsers.base import BaseParser, ParsedDocument, ParsedSection


class WordParser(BaseParser):
    """Parse .docx files, extracting structure via heading styles."""

    def can_handle(self, file_type: str) -> bool:
        return file_type.lower() in ("docx", "doc")

    def parse(self, file_path: Path) -> ParsedDocument:
        doc = DocxDocument(file_path)

        # Try to get title from document properties or first heading
        title = doc.core_properties.title or file_path.stem

        sections: list[ParsedSection] = []
        current_section: ParsedSection | None = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            heading_level = self._get_heading_level(style_name, para)

            if heading_level is not None:
                # This paragraph is a heading
                section = ParsedSection(
                    level=heading_level,
                    heading=text,
                    content=text,
                )
                sections.append(section)
                current_section = section

                # Use the first H1 as title if no title found
                if heading_level == 1 and title == file_path.stem:
                    title = text

            elif current_section is not None:
                # Append to current section
                current_section.content += "\n" + text
            else:
                # Before any heading, create a preamble section
                sections.append(ParsedSection(
                    level=0,
                    heading=title,
                    content=text,
                ))
                current_section = sections[-1]

        # Also try to extract table content
        for table in doc.tables:
            table_text_parts = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_text_parts.append(" | ".join(cells))
            if table_text_parts:
                table_content = "\n" + "\n".join(table_text_parts)
                if sections:
                    sections[-1].content += table_content
                else:
                    sections.append(ParsedSection(
                        level=0,
                        heading=title,
                        content=table_content,
                    ))

        return ParsedDocument(
            title=title,
            file_type="docx",
            sections=sections,
            metadata={
                "author": str(doc.core_properties.author or ""),
                "created": str(doc.core_properties.created or ""),
            },
        )

    def _get_heading_level(self, style_name: str, para) -> int | None:
        """Determine heading level from paragraph style."""
        style_lower = style_name.lower()

        # Standard heading styles
        if "heading" in style_lower:
            for i in range(1, 7):
                if f"heading {i}" in style_lower or f"heading{i}" in style_lower:
                    return i

        # Chinese common styles
        if "标题" in style_name:
            for i in range(1, 7):
                if str(i) in style_name:
                    return i
            return 1

        # Check outline level from XML
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            outline = pPr.find(qn('w:outlineLvl'))
            if outline is not None:
                level = int(outline.get(qn('w:val'), '0'))
                return level + 1

        return None
